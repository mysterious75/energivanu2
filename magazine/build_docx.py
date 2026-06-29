#!/usr/bin/env python3
"""Generate a professional DOCX version of the Energivanu magazine."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(SCRIPT_DIR, 'assets')
OUTPUT = os.path.join(SCRIPT_DIR, 'Energivanu_Insights_Magazine.docx')

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(200, 200, 200)

# Set narrow margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_dark_bg(paragraph):
    """Add dark background shading to paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '0A0A0A')
    pPr.append(shd)

def set_paragraph_bg(doc, hex_color='0A0A0A'):
    """Set document background."""
    background = OxmlElement('w:background')
    background.set(qn('w:color'), hex_color)
    doc.element.insert(0, background)

set_paragraph_bg(doc, '0D1B2A')

def add_heading_styled(text, level=1, color=RGBColor(0, 212, 255)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    add_dark_bg(h)
    return h

def add_body(text, bold=False, color=RGBColor(200, 200, 200), size=Pt(10)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size = size
    run.font.name = 'Calibri'
    run.bold = bold
    add_dark_bg(p)
    return p

def add_image_with_caption(filename, caption, width=Inches(6)):
    path = os.path.join(ASSETS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        add_dark_bg(p)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(100, 100, 100)
        r.font.name = 'Calibri'
        add_dark_bg(cap)

def add_table_styled(headers, rows, highlight=-1):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 212, 255)
                run.font.size = Pt(8)
                run.font.name = 'Calibri'
                run.bold = True
        # Dark bg
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '1A1A2E')
        tcPr.append(shd)
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(190, 190, 190) if i != highlight else RGBColor(255, 255, 255)
                    run.font.size = Pt(8)
                    run.font.name = 'Calibri'
                    run.bold = (i == highlight)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), '0F0F14' if i != highlight else '001E28')
            tcPr.append(shd)
    return table


# ==================== COVER ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ENERGIVANU INSIGHTS')
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(200, 16, 46)
r.font.name = 'Calibri'
r.bold = True
add_dark_bg(p)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('JUNE 2026 - ENERGY & AI')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(150, 150, 150)
r.font.name = 'Calibri'
add_dark_bg(p)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('EXCLUSIVE FEATURE')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0, 212, 255)
r.font.name = 'Calibri'
r.bold = True
add_dark_bg(p)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('The Open-Source Engine\nThat Could Save $47B\nin Data Center Power')
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(255, 255, 255)
r.font.name = 'Calibri'
r.bold = True
add_dark_bg(p)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('How a 613,000-parameter neural network is rewriting the rules of GPU data center energy management - and why ERCOT just made it essential.')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(150, 150, 150)
r.font.name = 'Calibri'
add_dark_bg(p)

# Stats
stats_text = '30% Grid Smoothing  |  59% Variance Reduction  |  <21s Response Time  |  100% Open Source'
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(stats_text)
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0, 255, 136)
r.font.name = 'Calibri'
r.bold = True
add_dark_bg(p)

doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
add_heading_styled('Inside This Issue', 1)
add_body('A deep dive into the technology, the market, and the mission that is reshaping data center energy.', color=RGBColor(150, 150, 150))

toc = [
    ('03', 'The $47 Billion Problem', "Why AI data centers are the new frontier of energy crisis."),
    ('04', 'Enter Energivanu', 'The open-source ML toolkit combining power prediction, battery dispatch, and phase staggering.'),
    ('05', 'Architecture Deep Dive', 'TCN + Attention, MPC, and the 15-feature input system.'),
    ('06', 'Training on 30 Lakh Rows', 'From 8,438% MAPE to 20.3% on Alibaba real telemetry.'),
    ('07', 'Verified Performance', 'Real hardware validation on Tesla P100. All metrics verified.'),
    ('08', 'The Competitive Edge', 'Comparison with Zeus, Emerald AI, Phaidra, and the landscape.'),
    ('09', 'Market & Opportunity', '$47B TAM by 2030. Where Energivanu fits.'),
    ('10', 'The Road Ahead', 'Production pilots, DCGM integration, real BESS hardware.'),
]

for num, title, desc in toc:
    p = doc.add_paragraph()
    r = p.add_run(f'{num}  ')
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0, 212, 255)
    r.font.name = 'Calibri'
    r.bold = True
    r = p.add_run(title)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.name = 'Calibri'
    r.bold = True
    add_dark_bg(p)
    p2 = doc.add_paragraph()
    r = p2.add_run(desc)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(120, 120, 120)
    r.font.name = 'Calibri'
    add_dark_bg(p2)

doc.add_page_break()

# ==================== THE PROBLEM ====================
add_heading_styled('The $47 Billion Problem', 1)
add_body("AI's insatiable appetite for power is creating a crisis that grid operators, data center operators, and regulators can no longer ignore.", color=RGBColor(150, 150, 150))

add_body("When ERCOT approved the Passive Controllable Load Resource framework on June 18, 2026, it was acknowledging a fundamental shift in how the electrical grid interacts with the largest power consumers on the planet.", bold=True, color=RGBColor(240, 240, 240), size=Pt(11))

add_body("The numbers are staggering. Global data center power demand is projected to reach 1,260 GW by 2030, with AI workloads consuming an estimated 820 GW of that total. The ERCOT grid alone faces a 410 GW large-load queue - and 87% of those applications are for data centers.")

add_body("The problem is not just raw power consumption. It is volatility. When a distributed AI training job triggers an All-Reduce synchronization across thousands of GPUs, the power draw can spike by megawatts in seconds.")

add_image_with_caption('market_opportunity.png', 'Global Data Center Power Demand Projection (2024-2030)')

add_heading_styled("ERCOT's Answer: PCLR", 2)
add_body("The PCLR framework offers data centers a deal: demonstrate load flexibility, get faster interconnection.")

add_table_styled(
    ['Requirement', 'Specification', 'Challenge'],
    [
        ['Dispatchability', 'Respond to SCED base points', 'Within 10 minutes'],
        ['Telemetry', 'ICCP communication with ERCOT', '4-second intervals'],
        ['Compliance', 'Stay within deadband', '+/- 5 MW tolerance'],
        ['Registration', 'QSE via RIOO', 'Form W by Jul 10/24, 2026'],
    ]
)

doc.add_page_break()

# ==================== ENTER ENERGIVANU ====================
add_heading_styled('Enter Energivanu', 1)
add_body('The open-source execution engine that fills the gap between grid signals and GPU clusters.', color=RGBColor(150, 150, 150))

add_body("Energivanu is what happens when you treat data center power optimization as a machine learning problem instead of a facilities management problem. The result is a 613,000-parameter neural network that can predict power spikes, dispatch battery storage, and stagger training phases - all in under 21 seconds.", bold=True, color=RGBColor(240, 240, 240), size=Pt(11))

add_body("GPU data centers do not have a power problem. They have a coordination problem. When a thousand GPUs start an All-Reduce operation simultaneously, the power spike is unpredictable. Traditional demand response systems cannot handle millisecond-scale volatility from compute workloads.")

add_heading_styled('Three Engines, One Pipeline', 2)

add_heading_styled('1. Power Prediction Engine', 3)
add_body('A Temporal Convolutional Network with Multi-Head Self-Attention that ingests 15 features across 30 timesteps and forecasts power consumption 10 steps ahead.')

add_heading_styled('2. Battery Dispatch Engine', 3)
add_body('A Model Predictive Controller that computes optimal charge/discharge trajectories for Battery Energy Storage Systems.')

add_heading_styled('3. Phase Coordination Engine', 3)
add_body('A scheduler that calculates optimal offsets for All-Reduce communication phases across GPU clusters, reducing aggregate grid volatility by up to 59%.')

add_image_with_caption('architecture.png', 'Energivanu System Architecture')

doc.add_page_break()

# ==================== ARCHITECTURE ====================
add_heading_styled('Architecture Deep Dive', 1)
add_body('Inside the neural network, the controller, and the features that make it work.', color=RGBColor(150, 150, 150))

add_heading_styled('The Neural Network: EnergivanuPEB', 2)
add_body('The core model uses a Temporal Convolutional Network (TCN) backbone with dilated causal convolutions, followed by 8-head self-attention.')

add_image_with_caption('feature_importance.png', '15-Feature Input Architecture')

add_heading_styled('Model Predictive Control', 2)
add_body('The MPC controller solves a constrained optimization problem at every decision step. Objective: minimize Sum[ Q*(P_grid - P_target)^2 + R*u^2 + S*(u_k - u_{k-1})^2 ]', bold=True)

add_body('Q = 100.0 (grid deviation)  |  R = 0.01 (battery wear)  |  S = 0.1 (ramp rate)')
add_body('Constraints: SOC 5-95%  |  Power -319.2 to +319.2 MW  |  Ramp 5 MW/min')

add_image_with_caption('response_timeline.png', 'End-to-End Response: Signal Parsing to Execution in Under 21 Seconds')

doc.add_page_break()

# ==================== TRAINING ====================
add_heading_styled('Training on 30 Lakh Rows', 1)
add_body('From 8,438% error to 20.3% - the iterative journey of building a production-grade prediction model.', color=RGBColor(150, 150, 150))

add_image_with_caption('training_progression.png', 'MAPE Improvement - 99.8% Error Reduction')

add_heading_styled('The Training Journey', 2)
add_table_styled(
    ['Version', 'Data Source', 'Rows', 'Params', 'Val Loss', 'MAPE'],
    [
        ['v1', 'MIT Supercloud', '14K', '338K', '0.0002', '8,438%'],
        ['v2', 'Alibaba (processed)', '3L', '338K', '88.0', '75.4%'],
        ['v3', 'Alibaba (full proc)', '50L', '338K', '34.59', '37.28%'],
        ['v4 *', 'Alibaba (raw sensor)', '30L', '613K', '5.95', '~20.3%'],
    ],
    highlight=3
)

add_heading_styled('Key Insight: Raw > Processed', 2)
add_body('Raw sensor data preserves the high-frequency power spikes and transient patterns that processed data smooths away. Combined with a larger model, this dropped MAPE from 37% to 20%.')

add_image_with_caption('alibaba_training.png', 'Training & Validation Loss Curve')

doc.add_page_break()

# ==================== VERIFIED PERFORMANCE ====================
add_heading_styled('Verified Performance', 1)
add_body('Real hardware. Real data. Real results.', color=RGBColor(150, 150, 150))

add_heading_styled('Key Metrics', 2)
add_table_styled(
    ['Metric', 'Value', 'Method'],
    [
        ['Grid Smoothing', '30.0%', 'MPCController, 30-step trace'],
        ['Peak Shaving', '10.5%', 'PeakShavingOptimizer, 24h TOU'],
        ['Phase Staggering', '59.0%', '4 clusters, seed=42'],
        ['Model Parameters', '613,612', 'TCN + Attention'],
        ['MAPE (Alibaba)', '~20.3%', '30L rows, 200 epochs'],
        ['PCLR Compliance', 'PASS', '5 MW deadband, 600s deadline'],
    ]
)

add_image_with_caption('bess_smoothing.png', 'Before & After BESS Optimization')

add_heading_styled('Real Hardware Validation', 2)
add_table_styled(
    ['Test', 'Status', 'Result'],
    [
        ['Production Telemetry', 'PASS', '26.3W avg idle, 35C'],
        ['MPC Controller', 'PASS', '30.0% grid smoothing'],
        ['BESS Physics', 'PASS', 'Stable SOC tracking'],
        ['Grid Integration', 'PASS', 'PCLR compliant'],
        ['Peak Shaving', 'PASS', '10.5% reduction'],
        ['Phase Staggering', 'PASS', '59.0% variance reduction'],
    ]
)

doc.add_page_break()

# ==================== COMPETITIVE ====================
add_heading_styled('The Competitive Edge', 1)
add_body('Energivanu is the only open-source project combining ML power prediction + BESS MPC + phase staggering.', color=RGBColor(150, 150, 150))

add_image_with_caption('competitive_radar.png', 'Competitive Capability Matrix')

add_table_styled(
    ['Project', 'Layer', 'BESS', 'GPU', 'Phase', 'Grid', 'Open'],
    [
        ['Energivanu', 'Cluster', 'Yes', 'Yes', 'Yes', 'Yes', 'AGPLv3'],
        ['Emerald AI', 'Grid', 'No', 'No', 'No', 'Yes', 'Proprietary'],
        ['Phaidra', 'Cooling', 'No', 'Partial', 'No', 'No', 'Proprietary'],
        ['FlexGen', 'Facility', 'Yes', 'No', 'No', 'Partial', 'Proprietary'],
        ['Zeus', 'GPU', 'No', 'Yes', 'No', 'No', 'Apache 2.0'],
        ['RADDiT', 'Cluster', 'No', 'Yes', 'No', 'Yes', 'Open'],
        ['GridPilot', 'Cluster', 'No', 'Yes', 'No', 'Yes', 'Research'],
    ],
    highlight=0
)

add_heading_styled('The Integration Advantage', 2)
add_body('No other project combines ML power prediction + BESS MPC + phase staggering in a single, deployable package.')

add_heading_styled('The Open-Source Advantage', 2)
add_body('Proprietary solutions require vendor lock-in and six-figure contracts. Energivanu is AGPLv3 - free to use, modify, and deploy.')

doc.add_page_break()

# ==================== MARKET ====================
add_heading_styled('Market & Opportunity', 1)
add_body('The convergence of AI scaling, grid constraints, and regulation creates a once-in-a-generation opportunity.', color=RGBColor(150, 150, 150))

add_heading_styled('The Numbers', 2)
add_table_styled(
    ['Metric', 'Value', 'Context'],
    [
        ['TAM by 2030', '$47B', 'DC Power Optimization'],
        ['ERCOT Queue', '410 GW', '87% Data Centers'],
        ['Global DC Power', '1,260 GW', 'Projected 2030'],
    ]
)

add_heading_styled('Why Now?', 2)
add_body('1. ERCOT PCLR Framework (June 18, 2026) - First major grid operator to create formal load flexibility for data centers.')
add_body('2. AI Compute Scaling - GPU clusters doubling every 12-18 months. H100s consume 700W each.')
add_body('3. Grid Capacity Limits - Transmission takes 3-5 years to build. DCs want 12-18 months.')

add_heading_styled('Revenue Model', 2)
add_table_styled(
    ['Stream', 'Description', 'Market'],
    [
        ['Commercial License', 'Proprietary deployment without AGPL', 'Enterprise DC operators'],
        ['Integration Services', 'Custom DCGM, BESS, grid integration', 'Colocation providers'],
        ['Retraining Services', 'Facility-specific model training', 'Hyperscalers'],
        ['Compliance Consulting', 'PCLR registration and support', 'Texas DC operators'],
    ]
)

doc.add_page_break()

# ==================== ROAD AHEAD ====================
add_heading_styled('The Road Ahead', 1)
add_body('From Kaggle notebooks to production data centers.', color=RGBColor(150, 150, 150))

add_heading_styled('Development Roadmap', 2)
add_table_styled(
    ['Timeline', 'Milestone', 'Description'],
    [
        ['Q3 2026', 'Production Pilot', '16-32 GPU validation at university or neocloud'],
        ['Q4 2026', 'DCGM Integration', 'Replace nvidia-smi with NVIDIA DCGM'],
        ['Q1 2027', 'Real BESS Hardware', 'Tesla Megapack Modbus client'],
        ['Q2 2027', 'OpenADR VTN', 'Real OpenADR test infrastructure'],
        ['H2 2027', 'Fleet Management', 'Multi-cluster aggregation API'],
    ]
)

add_heading_styled('The Vision', 2)
add_body('Every GPU data center should have an intelligent power management layer. Not as a luxury, but as fundamental infrastructure - as essential as networking or cooling.', bold=True, color=RGBColor(240, 240, 240), size=Pt(11))

add_body('The AI revolution is built on compute. Compute runs on power. And power, unlike compute, is bounded by physics, geography, and grid infrastructure.')

# Founder
p = doc.add_paragraph()
r = p.add_run('\nVed Kumar - Creator & Lead Developer')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0, 212, 255)
r.font.name = 'Calibri'
r.bold = True
add_dark_bg(p)

p = doc.add_paragraph()
r = p.add_run('Built from scratch on Kaggle free tier. Open-source advocate. Available for consulting, partnerships, and pilot deployments.')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(150, 150, 150)
r.font.name = 'Calibri'
add_dark_bg(p)

# CTA
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\nJoin the Revolution')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(200, 16, 46)
r.font.name = 'Calibri'
r.bold = True
add_dark_bg(p)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('github.com/mysterious75/energivanu2  |  @VEDKUMAR98  |  Open Source  |  AGPL-3.0')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(150, 150, 150)
r.font.name = 'Calibri'
add_dark_bg(p)

doc.save(OUTPUT)
size_mb = os.path.getsize(OUTPUT) / (1024 * 1024)
print(f"DOCX generated: {OUTPUT}")
print(f"Size: {size_mb:.2f} MB")
